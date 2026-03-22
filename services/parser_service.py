"""
Parser service for extracting text and information from CV files.
Supports PDF and DOCX formats.
"""
import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path

try:
    import fitz  # PyMuPDF — eng yaxshi matn chiqarish (Word/Canva PDF)
except ImportError:
    fitz = None

try:
    from pypdf import PdfReader as PypdfReader
except ImportError:
    PypdfReader = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class CVParserService:
    """
    Service class for parsing CV files and extracting structured information.
    """
    
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx']
    
    # Common skill keywords for basic extraction
    SKILL_KEYWORDS = {
        'programming': [
            'python', 'java', 'javascript', 'c++', 'c#', 'php', 'ruby', 'go', 
            'swift', 'kotlin', 'typescript', 'scala', 'rust', 'perl', 'r',
            'html', 'css', 'sql', 'nosql', 'mongodb', 'postgresql', 'mysql'
        ],
        'frameworks': [
            'django', 'flask', 'fastapi', 'react', 'angular', 'vue', 'node',
            'express', 'spring', 'laravel', 'rails', 'asp.net', 'symfony'
        ],
        'technologies': [
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
            'linux', 'apache', 'nginx', 'redis', 'elasticsearch', 'kafka'
        ],
        'methodologies': [
            'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'tdd', 'bdd'
        ]
    }
    
    def __init__(self):
        """Initialize the parser service."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are installed."""
        if fitz is None and PypdfReader is None and PyPDF2 is None:
            logger.warning("No PDF library (PyMuPDF / pypdf / PyPDF2). PDF parsing will fail.")

        if Document is None:
            logger.warning("python-docx not installed. DOCX parsing will not work.")
    
    def is_supported_file(self, file_path: str) -> bool:
        """
        Check if the file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            bool: True if file format is supported
        """
        return Path(file_path).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF using several backends (PyMuPDF → pypdf → PyPDF2).

        PyPDF2 alone often returns empty strings for PDFs exported from Word / Canva
        or with non-standard encodings; PyMuPDF usually fixes that.
        """
        path = str(file_path)
        errors: List[str] = []

        # 1) PyMuPDF
        if fitz is not None:
            try:
                doc = fitz.open(path)
                try:
                    parts = []
                    for i in range(doc.page_count):
                        parts.append(doc.load_page(i).get_text("text") or "")
                    text = "\n".join(parts).strip()
                    if text:
                        logger.info("PDF text via PyMuPDF: %d chars", len(text))
                        return text
                    errors.append("PyMuPDF: empty text (image-only / scanned PDF?)")
                finally:
                    doc.close()
            except Exception as e:
                err = f"PyMuPDF: {e}"
                errors.append(err)
                logger.warning(err)

        # 2) pypdf (maintained fork, better than old PyPDF2 in many cases)
        if PypdfReader is not None:
            try:
                reader = PypdfReader(path, strict=False)
                chunks = []
                for page in reader.pages:
                    chunks.append(page.extract_text() or "")
                text = "\n".join(chunks).strip()
                if text:
                    logger.info("PDF text via pypdf: %d chars", len(text))
                    return text
                errors.append("pypdf: empty text")
            except Exception as e:
                err = f"pypdf: {e}"
                errors.append(err)
                logger.warning(err)

        # 3) PyPDF2 (last resort)
        if PyPDF2 is not None:
            try:
                with open(path, "rb") as fh:
                    pdf_reader = PyPDF2.PdfReader(fh)
                    text = "\n".join(
                        (p.extract_text() or "") for p in pdf_reader.pages
                    ).strip()
                    if text:
                        logger.info("PDF text via PyPDF2: %d chars", len(text))
                        return text
                errors.append("PyPDF2: empty text")
            except Exception as e:
                err = f"PyPDF2: {e}"
                errors.append(err)
                logger.warning(err)

        if not errors:
            raise ImportError(
                "Install a PDF library: pip install PyMuPDF pypdf PyPDF2"
            )

        raise ValueError(
            "Could not extract text from PDF. "
            "If this is a scanned (image) PDF, export as DOCX or use OCR. "
            f"Attempts: {'; '.join(errors)}"
        )
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ImportError: If python-docx is not installed
            Exception: If file cannot be processed
        """
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a CV file (PDF or DOCX).
        
        Args:
            file_path: Path to the CV file
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If file cannot be processed
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def extract_email(self, text: str) -> Optional[str]:
        """
        Extract email address from text.
        
        Args:
            text: Text content to search
            
        Returns:
            str or None: First email address found
        """
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """
        Extract phone number from text.
        
        Args:
            text: Text content to search
            
        Returns:
            str or None: First phone number found
        """
        # Simple phone pattern - can be improved
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\(\d{3}\)\s?\d{3}-\d{4}',
            r'\d{3}-\d{3}-\d{4}'
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        
        return None
    
    def _name_from_filename(self, filename: str) -> Optional[str]:
        """
        Guess candidate name from upload filename, e.g. ``AzizbekGulomov_CV.pdf``.
        """
        stem = Path(filename).stem
        stem = re.sub(r"(?i)[_\-]?(cv|resume|curriculum\s*vitae)[_\-]?", " ", stem)
        stem = stem.strip(" _-.")
        if not stem:
            return None
        # camelCase / PascalCase -> words
        spaced = re.sub(r"([a-z])([A-Z])", r"\1 \2", stem)
        spaced = re.sub(r"[_\-.]+", " ", spaced)
        words = re.findall(r"[A-Za-zÀ-ÿ']+", spaced)
        if 2 <= len(words) <= 6:
            return " ".join(w.capitalize() for w in words)
        if len(words) == 1 and 2 < len(words[0]) < 40:
            return words[0][:1].upper() + words[0][1:].lower()
        return None
    
    def extract_candidate_name(self, text: str, filename: str) -> Optional[str]:
        """
        Best-effort name: filename first, then first plausible header line in text.
        """
        from_file = self._name_from_filename(filename)
        if from_file:
            return from_file[:200]
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for line in lines[:8]:
            if "@" in line or len(line) > 70:
                continue
            if re.search(r"\d{3}[-.\s]?\d{2}", line):
                continue
            # 2–5 title-case / letter words (Latin + common UTF-8 letters)
            if re.match(
                r"^[\w'\-]+\s+[\w'\-]+(?:\s+[\w'\-]+){0,3}$",
                line,
                re.UNICODE,
            ) and not line.lower().startswith("http"):
                return line[:200]
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from CV text using keyword matching.
        
        Args:
            text: CV text content
            
        Returns:
            List[str]: List of identified skills
        """
        text_lower = text.lower()
        found_skills = []
        
        # Search for skills in all categories
        for category, skills in self.SKILL_KEYWORDS.items():
            for skill in skills:
                # Use word boundaries to avoid partial matches
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.append(skill.title())
        
        # Remove duplicates and sort
        return sorted(list(set(found_skills)))
    
    def extract_experience_years(self, text: str) -> Optional[int]:
        """
        Extract years of experience from CV text.
        
        Args:
            text: CV text content
            
        Returns:
            int or None: Number of years of experience found
        """
        # Patterns to match experience mentions
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*years?\s*in\s*(?:the\s*)?field',
            r'experience:\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*professional\s*experience'
        ]
        
        text_lower = text.lower()
        years = []
        
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years.append(int(match))
                except ValueError:
                    continue
        
        # Return the maximum years found (most likely to be total experience)
        return max(years) if years else None
    
    def extract_education(self, text: str) -> str:
        """
        Extract education information from CV text.
        
        Args:
            text: CV text content
            
        Returns:
            str: Education information found
        """
        education_keywords = [
            'education', 'academic', 'qualification', 'degree', 'university',
            'college', 'bachelor', 'master', 'phd', 'doctorate', 'diploma',
            'certificate', 'coursework'
        ]
        
        lines = text.split('\n')
        education_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in education_keywords):
                # Include this line and potentially next few lines
                education_lines.append(line.strip())
        
        return '\n'.join(education_lines) if education_lines else ""
    
    def parse_cv(self, file_path: str) -> Dict:
        """
        Parse a CV file and extract structured information.
        
        Args:
            file_path: Path to the CV file
            
        Returns:
            dict: Structured CV information
        """
        try:
            basename = Path(file_path).name
            # Extract text from file
            text = self.extract_text(file_path)
            
            if not text.strip():
                raise ValueError("No text could be extracted from the file")
            
            # Extract structured information
            result = {
                'extracted_text': text,
                'name': self.extract_candidate_name(text, basename),
                'email': self.extract_email(text),
                'phone': self.extract_phone(text),
                'skills': self.extract_skills(text),
                'experience_years': self.extract_experience_years(text),
                'education': self.extract_education(text),
                'parsing_success': True,
                'error_message': None
            }
            
            logger.info(f"Successfully parsed CV: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse CV {file_path}: {str(e)}")
            return {
                'extracted_text': "",
                'name': None,
                'email': None,
                'phone': None,
                'skills': [],
                'experience_years': None,
                'education': "",
                'parsing_success': False,
                'error_message': str(e)
            }
    
    def get_file_info(self, file_path: str) -> Dict:
        """
        Get basic information about a CV file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            dict: File information
        """
        path_obj = Path(file_path)
        
        return {
            'filename': path_obj.name,
            'extension': path_obj.suffix.lower(),
            'size_bytes': path_obj.stat().st_size if path_obj.exists() else 0,
            'is_supported': self.is_supported_file(file_path)
        }